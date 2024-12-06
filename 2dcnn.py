import os
import random
import tensorflow as tf
from tensorflow.keras.preprocessing.image import ImageDataGenerator # type: ignore
from tensorflow.keras.models import Sequential # type: ignore
from tensorflow.keras.layers import Conv2D, MaxPooling2D, Flatten, Dense, Dropout # type: ignore
from tensorflow.keras.callbacks import EarlyStopping # type: ignore
import matplotlib.pyplot as plt
from sklearn.metrics import classification_report, confusion_matrix, ConfusionMatrixDisplay
from docx import Document # type: ignore
from fpdf import FPDF # type: ignore
import numpy as np

# Set seeds for reproducibility
seed_value = 42
os.environ['PYTHONHASHSEED'] = str(seed_value)
random.seed(seed_value)
np.random.seed(seed_value)
tf.random.set_seed(seed_value)

# Check if TensorFlow can access the GPU
gpus = tf.config.list_physical_devices('GPU')
if gpus:
    print("Using GPU")
    # Optionally, set memory growth to prevent over-allocation
    for gpu in gpus:
        tf.config.experimental.set_memory_growth(gpu, True)
else:
    print("No GPU found. Using CPU.")



# Paths
dataset_dir = './pneumonia-dataset'
train_dir = os.path.join(dataset_dir, 'train')
val_dir = os.path.join(dataset_dir, 'val')
test_dir = os.path.join(dataset_dir, 'test')

# Define CNN model
def build_model():
    model = Sequential([
        Conv2D(32, (3, 3), activation='relu', input_shape=(150, 150, 3)),
        MaxPooling2D(pool_size=(2, 2)),
        Conv2D(64, (3, 3), activation='relu'),
        MaxPooling2D(pool_size=(2, 2)),
        Conv2D(128, (3, 3), activation='relu'),
        MaxPooling2D(pool_size=(2, 2)),
        Flatten(),
        Dense(128, activation='relu'),
        Dropout(0.5),
        Dense(1, activation='sigmoid')  # Binary classification
    ])
    model.compile(optimizer='adam',
                  loss='binary_crossentropy',
                  metrics=['accuracy'])
    return model

# Data Generators
simple_datagen = ImageDataGenerator(rescale=1.0/255.0)
augmented_datagen = ImageDataGenerator(
    rescale=1.0/255.0,
    rotation_range=20,
    width_shift_range=0.2,
    height_shift_range=0.2,
    shear_range=0.2,
    zoom_range=0.2,
    horizontal_flip=True
)

# Train Without Augmentation
print("Training without data augmentation...")
simple_train_generator = simple_datagen.flow_from_directory(
    train_dir, target_size=(150, 150), batch_size=32, class_mode='binary'
)
simple_val_generator = simple_datagen.flow_from_directory(
    val_dir, target_size=(150, 150), batch_size=32, class_mode='binary'
)

simple_model = build_model()
early_stopping = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)

history_simple = simple_model.fit(
    simple_train_generator,
    steps_per_epoch=simple_train_generator.samples // simple_train_generator.batch_size,
    epochs=20,
    validation_data=simple_val_generator,
    validation_steps=simple_val_generator.samples // simple_val_generator.batch_size,
    callbacks=[early_stopping]
)

# Evaluate Without Augmentation
test_generator = simple_datagen.flow_from_directory(
    test_dir, target_size=(150, 150), batch_size=32, class_mode='binary', shuffle=False
)
simple_loss, simple_accuracy = simple_model.evaluate(test_generator)
print(f"Without Augmentation - Loss: {simple_loss}, Accuracy: {simple_accuracy}")

# Train With Augmentation
print("\nTraining with data augmentation...")
train_generator = augmented_datagen.flow_from_directory(
    train_dir, target_size=(150, 150), batch_size=32, class_mode='binary'
)
val_generator = simple_datagen.flow_from_directory(
    val_dir, target_size=(150, 150), batch_size=32, class_mode='binary'
)

augmented_model = build_model()
history_augmented = augmented_model.fit(
    train_generator,
    steps_per_epoch=train_generator.samples // train_generator.batch_size,
    epochs=20,
    validation_data=val_generator,
    validation_steps=val_generator.samples // val_generator.batch_size,
    callbacks=[early_stopping]
)

# Evaluate With Augmentation
augmented_loss, augmented_accuracy = augmented_model.evaluate(test_generator)
print(f"With Augmentation - Loss: {augmented_loss}, Accuracy: {augmented_accuracy}")

# Classification Reports
simple_y_pred = (simple_model.predict(test_generator) > 0.5).astype("int32")
augmented_y_pred = (augmented_model.predict(test_generator) > 0.5).astype("int32")
y_true = test_generator.classes

simple_cr = classification_report(y_true, simple_y_pred, target_names=test_generator.class_indices.keys())
augmented_cr = classification_report(y_true, augmented_y_pred, target_names=test_generator.class_indices.keys())
print("Without Augmentation Classification Report:\n", simple_cr)
print("With Augmentation Classification Report:\n", augmented_cr)

# Confusion Matrices
fig, ax = plt.subplots(1, 2, figsize=(12, 6))
simple_cm = confusion_matrix(y_true, simple_y_pred)
augmented_cm = confusion_matrix(y_true, augmented_y_pred)

ConfusionMatrixDisplay(simple_cm, display_labels=test_generator.class_indices.keys()).plot(ax=ax[0], cmap='Blues')
ax[0].set_title('Without Augmentation')
ConfusionMatrixDisplay(augmented_cm, display_labels=test_generator.class_indices.keys()).plot(ax=ax[1], cmap='Blues')
ax[1].set_title('With Augmentation')

plt.tight_layout()
plt.savefig('confusion_matrices.png')
plt.close()

# Plot Comparison
plt.figure(figsize=(12, 6))

# Loss Comparison
plt.subplot(1, 2, 1)
plt.plot(history_simple.history['val_loss'], label='Without Augmentation')
plt.plot(history_augmented.history['val_loss'], label='With Augmentation')
plt.title('Validation Loss Comparison')
plt.xlabel('Epochs')
plt.ylabel('Loss')
plt.legend()

# Accuracy Comparison
plt.subplot(1, 2, 2)
plt.plot(history_simple.history['val_accuracy'], label='Without Augmentation')
plt.plot(history_augmented.history['val_accuracy'], label='With Augmentation')
plt.title('Validation Accuracy Comparison')
plt.xlabel('Epochs')
plt.ylabel('Accuracy')
plt.legend()

plt.tight_layout()
plt.savefig('comparison_plot.png')
plt.close()

# Bar Graph for Accuracy and Loss
metrics = {
    "Without Augmentation": {
        "Accuracy": simple_accuracy,
        "Loss": simple_loss
    },
    "With Augmentation": {
        "Accuracy": augmented_accuracy,
        "Loss": augmented_loss
    }
}

fig, ax = plt.subplots(1, 2, figsize=(12, 6))

# Accuracy Bar Graph
ax[0].bar(metrics.keys(), [metrics["Without Augmentation"]["Accuracy"], metrics["With Augmentation"]["Accuracy"]], color=['blue', 'orange'])
ax[0].set_title('Accuracy Comparison')
ax[0].set_ylabel('Accuracy')
for i, v in enumerate([metrics["Without Augmentation"]["Accuracy"], metrics["With Augmentation"]["Accuracy"]]):
    ax[0].text(i, v + 0.02, f"{v:.2f}", ha='center', fontsize=12)

# Loss Bar Graph
ax[1].bar(metrics.keys(), [metrics["Without Augmentation"]["Loss"], metrics["With Augmentation"]["Loss"]], color=['blue', 'orange'])
ax[1].set_title('Loss Comparison')
ax[1].set_ylabel('Loss')
for i, v in enumerate([metrics["Without Augmentation"]["Loss"], metrics["With Augmentation"]["Loss"]]):
    ax[1].text(i, v + 0.02, f"{v:.2f}", ha='center', fontsize=12)

plt.tight_layout()
plt.savefig('bar_comparison_plot.png')
plt.close()

# Generate Reports
def generate_reports():
    # Word Report
    doc = Document()
    doc.add_heading('Model Performance Report', level=1)

    doc.add_heading('Comparison of Loss and Accuracy', level=2)
    doc.add_picture('comparison_plot.png')

    doc.add_heading('Accuracy and Loss Comparison (Bar Graph)', level=2)
    doc.add_picture('bar_comparison_plot.png')

    doc.add_heading('Confusion Matrices', level=2)
    doc.add_picture('confusion_matrices.png')

    doc.add_heading('Without Augmentation - Classification Report', level=2)
    doc.add_paragraph(simple_cr)

    doc.add_heading('With Augmentation - Classification Report', level=2)
    doc.add_paragraph(augmented_cr)

    doc.save('performance_report.docx')

    # PDF Report
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)

    pdf.multi_cell(0, 10, 'Model Performance Report\n')
    pdf.image('comparison_plot.png', x=10, y=30, w=180)
    pdf.add_page()
    pdf.image('bar_comparison_plot.png', x=10, y=30, w=180)
    pdf.add_page()
    pdf.image('confusion_matrices.png', x=10, y=10, w=180)

    pdf.set_y(200)
    pdf.multi_cell(0, 10, 'Without Augmentation - Classification Report\n' + simple_cr)
    pdf.add_page()
    pdf.multi_cell(0, 10, 'With Augmentation - Classification Report\n' + augmented_cr)

    pdf.output('performance_report.pdf')

generate_reports()
